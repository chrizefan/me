#!/usr/bin/env python3
"""Update Chris_Stefan_CV.docx to match cv.html content."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
title_style = doc.styles["Title"]
heading_style = doc.styles["Heading 1"]
normal_style = doc.styles["Normal"]

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Chris Stefan")
run.bold = True
run.font.size = Pt(18)
run.font.name = "Inter"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Senior AI Engineer  |  Production Agentic AI  |  Remote-Ready  |  "
    "Open to Full-Time, Contract & Consulting"
)
run.font.size = Pt(10)
run.font.name = "Inter"
run.font.color.rgb = None  # default

doc.add_paragraph()

# Contact (compact)
contact = [
    "chris.stefan@proton.me",
    "+1 514.710.9601",
    "linkedin.com/in/chris-stefan",
    "github.com/chrizefan",
    "chrizefan.github.io/me",
    "EU Citizen · Relocating to Europe",
]
for line in contact:
    p = doc.add_paragraph(line, style="Normal")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "Inter"

doc.add_paragraph()

# Summary
p = doc.add_paragraph("SUMMARY", style="Heading 2")
p.runs[0].font.size = Pt(10)
p.runs[0].font.bold = True
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(6)

summary = (
    "AI Engineer with 5 years building production AI at the intersection of "
    "artificial intelligence and institutional finance. Lead developer & technical "
    "owner of PSP Investments' flagship agentic AI platform — 300 total users, 100 "
    "active & recurring, platform showcased at the Databricks Data & AI Summit. "
    "Full-stack ownership: architecture, SDK design, UI, onboarding, and "
    "stakeholder management across all asset classes. Seeking senior, lead, or "
    "consulting roles to build high-impact AI products."
)
p = doc.add_paragraph(summary, style="Normal")
p.paragraph_format.space_after = Pt(12)
for run in p.runs:
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = "Inter"

# Experience
p = doc.add_paragraph("EXPERIENCE", style="Heading 2")
p.runs[0].font.size = Pt(10)
p.runs[0].font.bold = True
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(6)

experiences = [
    {
        "title": "Senior AI Engineer, AlphaScience — PSP Investments",
        "date": "Sep 2022 – Present",
        "bullets": [
            "Sole engineer & full-stack owner of PSP's agentic AI platform — from initial POC through system architecture, design documentation, UI/UX, feature development, SDK packaging & distribution, PR reviews, and ongoing maintenance; 100+ active analysts in daily workflows, cutting research cycles ~10× and supporting 20–30% broader name coverage.",
            "AlphaScience SDK: wraps Databricks (Vector Search, SQL Warehouses, Genie Agents) & OpenAI (Responses + Files API), auto-auth via Databricks CLI & Azure Key Vault — composable agent/workflow base-class framework, proprietary tracing, inter-agent persistence, pre-built financial utility agents & skills; one-liner to spin up an agent or query a vector index",
            "Virtual Analyst Platform: custom chatbot UI with step-level tracing & observability (reasoning trace, web search, RAG retrieval, parallel tool execution, subagents), source traceability, structured data persistence layer with sandboxed analysis & live plotting, long-running workflow launcher & scheduled task subscriptions",
            "Living roadmap: weekly-evolving priorities with rapid integration of bleeding-edge AI developments into production — each feature independently scoped, designed as an SDK tool, and surfaced through the platform from concept to deployment",
            "Cost impact: in-house platform (~$1M/year all-in: data, infrastructure & engineering) displaces enterprise tooling priced at ~$20K/analyst/year (e.g., AlphaSense: $1.2M/year for 60 licenses) — at 100+ active users, avoided spend exceeds platform cost and scales with every analyst onboarded",
            "Leadership: platform selected for live demo at Databricks Data & AI Summit; mentored interns to ship production features independently; led onboarding for 100+ analysts; direct collaboration with C-suite across all asset classes",
        ],
        "pills": ["Agentic AI", "RAG", "Orchestration", "Azure AI", "Evals", "Observability"],
    },
    {
        "title": "Analyst, Digital Innovation — PSP Investments",
        "date": "Sep 2021 – Aug 2022 · Montreal, QC",
        "bullets": [
            "Co-developed Earnings Prediction Engine (ML beat/miss pipeline, point-in-time data, Power BI) and quantitative equity strategy for $8B active portfolio (+125 bps/year); built alt-data pipeline with data profiling, cleaning & feature engineering for equity factor research; delivered POC analyses to drive stakeholder buy-in across asset class teams",
        ],
        "pills": ["Quant Strategy", "Alt Data", "ML Pipelines", "Power BI", "scikit-learn"],
    },
    {
        "title": "Intern, Information Security — PSP Investments",
        "date": "May–Aug 2021",
        "bullets": [
            "Enterprise risk framework implementation; automated risk assessment repository updates, improving data currency and reporting efficiency",
        ],
        "pills": ["Cybersecurity", "Risk Analysis"],
    },
    {
        "title": "Business Analyst — Lunch à Porter",
        "date": "Jun 2020 – Jun 2021",
        "bullets": [
            "SEO & A/B testing drove 2× online conversion rate; inventory analysis for data-informed merchandising decisions",
        ],
        "pills": ["SEO", "A/B Testing", "Shopify", "Lightspeed", "Analytics"],
    },
]

for exp in experiences:
    p = doc.add_paragraph()
    run = p.add_run(exp["title"])
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Inter"
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph(exp["date"])
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "Inter"

    for bullet in exp["bullets"]:
        p = doc.add_paragraph(bullet, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = "Inter"

    pills = "  ".join(exp["pills"])
    p = doc.add_paragraph(pills)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.name = "Inter"

# Personal Project
p = doc.add_paragraph("PERSONAL PROJECT", style="Heading 2")
p.runs[0].font.size = Pt(10)
p.runs[0].font.bold = True
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph("digithings.ai · github.com/digithings-ai")
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(4)

project_desc = (
    "Fully modular multi-agent AI toolkit. LangGraph orchestration, LiteLLM universal "
    "LLM wrapper, pluggable RAG backends, built-in evals, Docker deployment, MCP & "
    "Open WebUI integration, OpenClaw layer for autonomous agent deployment. "
    'Direction: "hedge fund in a box" — merging agentic AI with quantitative finance.'
)
p = doc.add_paragraph(project_desc)
p.paragraph_format.space_after = Pt(4)
for run in p.runs:
    run.font.size = Pt(9)
    run.font.name = "Inter"

project_pills = "LangGraph  LiteLLM  Docker  MCP  Open WebUI  Azure AI  Ollama  OpenClaw"
p = doc.add_paragraph(project_pills)
for run in p.runs:
    run.font.size = Pt(8)
    run.font.name = "Inter"

doc.add_paragraph()

# Skills
p = doc.add_paragraph("SKILLS", style="Heading 2")
p.runs[0].font.size = Pt(10)
p.runs[0].font.bold = True
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(6)

skills = {
    "Gen AI & LLMs": "OpenAI API, Gemini API, Claude API, Hugging Face, Prompt Engineering, Vector DBs",
    "Data & Cloud": "Azure, GCP, Databricks, Spark, Delta Lake, PostgreSQL, SQL, Supabase, MLOps, CI/CD",
    "Programming": "Python, JavaScript, HTML, CSS, FastAPI, Pydantic, REST APIs, SDK Design, Git",
    "Visualization": "Plotly, Dash, Matplotlib, ECharts, Mermaid, Tableau",
    "AI Dev Tools": "Cursor, Antigravity, GitHub Copilot, Claude Code, Bolt.new",
    "Quant Trading": "QuantConnect, IBKR API, Nautilus Trader, TA-Lib",
    "Data": "LSEG, Worldscope, IBES, SEC Filings, Broker Research, Transcripts",
}
for cat, items in skills.items():
    p = doc.add_paragraph()
    run = p.add_run(cat + ": ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Inter"
    run = p.add_run(items)
    run.font.size = Pt(9)
    run.font.name = "Inter"
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# Education
p = doc.add_paragraph("EDUCATION", style="Heading 2")
p.runs[0].font.size = Pt(10)
p.runs[0].font.bold = True
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(6)

edu = [
    ("McGill University", "B.Com Finance & Business Analytics", "Desautels · 2019–2022", "Montreal, QC"),
    ("Champlain College", "DEC Computer Science & Mathematics", "2017–2019 · Honor Roll", "Montreal, QC"),
]
for school, degree, meta, location in edu:
    p = doc.add_paragraph()
    run = p.add_run(school)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Inter"
    p.add_run(f"\n{degree}\n{meta}\n{location}")
    for run in p.runs[1:]:
        run.font.size = Pt(9)
        run.font.name = "Inter"
    p.paragraph_format.space_after = Pt(6)

# Certifications
p = doc.add_paragraph("CERTIFICATIONS", style="Heading 2")
p.runs[0].font.size = Pt(10)
p.runs[0].font.bold = True
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run("Azure AI Fundamentals")
run.bold = True
run.font.size = Pt(9)
run.font.name = "Inter"
p.add_run("\nMicrosoft · Sep 2023")
for run in p.runs[1:]:
    run.font.size = Pt(9)
    run.font.name = "Inter"

doc.add_paragraph()

# Languages
p = doc.add_paragraph("LANGUAGES", style="Heading 2")
p.runs[0].font.size = Pt(10)
p.runs[0].font.bold = True
p.runs[0].font.name = "Inter"
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph("English — Fluent\nFrench — Fluent\nSpanish — Basic\nItalian — Basic")
for run in p.runs:
    run.font.size = Pt(9)
    run.font.name = "Inter"

doc.save("Chris_Stefan_CV.docx")
print("Updated Chris_Stefan_CV.docx")
